#!/usr/bin/python3

import sys, getopt
import numpy as np
from math import *
import subprocess
import os
import re # regexp package
import shutil
import tempfile

usage = '''A timing script to generate perf data and plot for major fft 1D/2D/3D cases

Usage:
\talltime.py
\t\t-b          Specify dload executable(optional)
\t\t-i          Specify test libraries for dload or test executables
\t\t              for regular test(appendable)
\t\t-o          Specify output directories for raw data
\t\t              appendable; defaults to "dir0", "dir1", etc.
\t\t-l          Specify labels for runs
\t\t              appendable; defaults to "dir0", "dir1", etc.
\t\t-w          output directory for graphs and final document
\t\t-S          plot speedup (default: 1, disabled: 0)
\t\t-t          data type: time (default) or gflops or roofline
\t\t-y          secondary acix type: none or gflops
\t\t-s          short run
\t\t-T          do not perform FFTs; just generate document
\t\t-f          document format: pdf (default) or docx
\t\t-g          generate graphs via Asymptote: 0(default) or 1
\t\t-d          device number (default: 0)
\t\t-N          Number of samples (default: 10)
\t\t-D          dims to test. default: 1,2,3
\t\t-R          runtype: report benchmark or efficiency
'''

def nextpow(val, radix):
    x = 1
    while(x <= val):
        x *= radix
    return x

# A class for generating data for figures.
class rundata:
    def __init__(self, label,
                 dimension, minsize, maxsize, nbatch, radix, ratio, ffttype,
                 direction, inplace):
        self.dimension = dimension
        self.minsize = minsize
        self.maxsize = maxsize
        self.nbatch = nbatch
        self.radix = radix
        self.ratio = ratio
        self.ffttype = ffttype
        self.precision = "double"
        self.inplace = inplace
        self.direction = direction
        self.label = label

    def outfilename(self, odir):
        outfile = ""
        outfile += "radix" + str(self.radix)
        outfile += "_dim" + str(self.dimension)
        outfile += "_" + self.precision
        outfile += "_n" + str(self.nbatch)
        if self.direction == 1:
            outfile += "_inv"
        if self.dimension > 1:
            outfile += "_ratio" + "_" + str(self.ratio[0])
        if self.dimension > 2:
            outfile += "_" + str(self.ratio[1])
        outfile += "_" + self.ffttype
        if self.inplace:
            outfile += "_inplace"
        else:
            outfile += "_outofplace"
        outfile += ".dat"
        outfile = os.path.join(odir, outfile)
        return outfile
        
    def runcmd(self, nsample, inlist, outdirlist, dloadexe):
        cmd = [os.path.join(sys.path[0],"timing.py")]

        if dloadexe == None:
            # When not using dload, we just have one input and output dir.
            cmd.append("-w")
            cmd.append(os.path.abspath(inlist[0]))
            cmd.append("-o")
            cmd.append(self.outfilename(outdirlist[0]))
        else:
            cmd.append("-w")
            cmd.append(dloadexe)
            for indir in inlist:
                cmd.append("-i")
                cmd.append(indir)
            for outdir in outdirlist:
                cmd.append("-o")
                cmd.append(self.outfilename(outdir))
        
        cmd.append("-N")
        cmd.append(str(nsample))
        
        cmd.append("-b")
        cmd.append(str(self.nbatch))
        
        cmd.append("-x")
        cmd.append(str(self.minsize))
        cmd.append("-X")
        cmd.append(str(self.maxsize))

        if self.dimension > 1:
            cmd.append("-y")
            cmd.append(str(self.minsize * self.ratio[0]))
            cmd.append("-Y")
            cmd.append(str(self.maxsize * self.ratio[0]))

        if self.dimension > 2:
            cmd.append("-z")
            cmd.append(str(self.minsize * self.ratio[1]))
            cmd.append("-Z")
            cmd.append(str(self.maxsize * self.ratio[1]))

        cmd.append("-r")
        cmd.append(str(self.radix))

        cmd.append("-D")
        cmd.append(str(self.direction))
        
        cmd.append("-d")
        cmd.append(str(self.dimension))

        cmd.append("-f")
        cmd.append(self.precision)
        
        if self.ffttype == "r2c":
            cmd.append("-R")
            
        return cmd

    def executerun(self, nsample, inlist, outdirlist, dloadexe):
        fout = tempfile.TemporaryFile(mode="w+")
        ferr = tempfile.TemporaryFile(mode="w+")

        if dloadexe != None:
            cmd = self.runcmd(nsample, inlist, outdirlist, dloadexe)
            print(" ".join(cmd))
            proc = subprocess.Popen(cmd,
                                    stdout=fout, stderr=ferr,
                                    env=os.environ.copy())
            
            # FIXME: copy log to multiple outputs?
            
            proc.wait()
            rc = proc.returncode
            if rc != 0:
                print("****fail****")
            
        else:
            for idx in range(min(len(inlist), len(outdirlist))):
                print(idx, ":", inlist[idx], "->", outdirlist[idx], flush=True)
                cmd = self.runcmd(nsample, [inlist[idx]], [outdirlist[idx]], None)
                print(" ".join(cmd))
                proc = subprocess.Popen(cmd,
                                        stdout=fout, stderr=ferr,
                                        env=os.environ.copy())
                proc.wait()
                rc = proc.returncode
                if rc != 0:
                    print("****fail****")

        return 0


# Figure class, which contains runs and provides commands to generate figures.
class figure:
    def __init__(self, name, caption):
        self.name = name
        self.runs = []
        self.caption = caption
    
    def inputfiles(self, outdirlist):
        import os
        files = []
        for run in self.runs:
            for outdir in outdirlist:
                files.append(run.outfilename(outdir))
        print(files)
        return files

    def labels(self, labellist):
        labels = []
        for run in self.runs:
            for label in labellist:
                labels.append(label + run.label)
        return labels
    
    def filename(self, outdir, docformat):
        outfigure = self.name
        outfigure += ".pdf"
        # if docformat == "pdf":
        #     outfigure += ".pdf"
        # if docformat == "docx":
        #     outfigure += ".png"
        return os.path.join(outdir, outfigure)
        
    def asycmd(self, docdir, outdirlist, labellist, docformat, datatype, ncompare, secondtype, just1dc2crad2):
        asycmd = ["asy"]
        
        asycmd.append("-f")
        asycmd.append("pdf")
        # if docformat == "pdf":
        #     asycmd.append("-f")
        #     asycmd.append("pdf")
        # if docformat == "docx":
        #     asycmd.append("-f")
        #     asycmd.append("png")
        #     asycmd.append("-render")
        #     asycmd.append("8")
        asycmd.append(os.path.join(sys.path[0],"datagraphs.asy"))

               
        asycmd.append("-u")
        inputfiles = self.inputfiles(outdirlist)
        asycmd.append('filenames="' + ",".join(inputfiles) + '"')

        asycmd.append("-u")
        labels = self.labels(labellist)
        asycmd.append('legendlist="' + ",".join(labels) + '"')

        asycmd.append("-u")
        asycmd.append('speedup=' + str(ncompare))

        if just1dc2crad2 :
            asycmd.append("-u")
            asycmd.append('just1dc2crad2=true')
            
        if secondtype == "gflops":
            asycmd.append("-u")
            asycmd.append('secondarygflops=true')

        if datatype == "gflops":
            asycmd.append("-u")
            asycmd.append('primaryaxis="gflops"')

        if datatype == "roofline":
            asycmd.append("-u")
            asycmd.append('primaryaxis="roofline"')
            # roofline on multiple devices doesn't really make sense; just use the first device
            with open(os.path.join(outdirlist[0], "gpuid.txt"), "r") as f:
                gpuid = f.read()
                asycmd.append("-u")
                asycmd.append('gpuid="' + gpuid.strip() + '"')
                        
        if len(self.runs) > 0:
            asycmd.append("-u")
            asycmd.append('batchsize=' + str(self.runs[0].nbatch))
            asycmd.append("-u")
            asycmd.append('problemdim=' + str(self.runs[0].dimension))
            asycmd.append("-u")
            val = 1
            for rat in self.runs[0].ratio:
                val *= rat
            asycmd.append('problemratio=' + str(val))
            asycmd.append("-u")
            if self.runs[0].ffttype == "r2c":
                asycmd.append("realcomplex=true")
            else:
                asycmd.append("realcomplex=false")
            
            
        asycmd.append("-o")
        asycmd.append(self.filename(docdir, docformat) )
                    
        return asycmd

    def executeasy(self, docdir, outdirs, labellist, docformat, datatype, ncompare, secondtype,
                   just1dc2crad2):
        fout = tempfile.TemporaryFile(mode="w+")
        ferr = tempfile.TemporaryFile(mode="w+")
        asyproc = subprocess.Popen(self.asycmd(docdir, outdirs, labellist, 
                                               docformat, datatype, ncompare, secondtype,
                                               just1dc2crad2),
                                   stdout=fout, stderr=ferr, env=os.environ.copy(),
                                   cwd = sys.path[0])
        asyproc.wait()
        asyrc = asyproc.returncode
        if asyrc != 0:
            print("****asy fail****")
            fout.seek(0)
            cout = fout.read()
            print(cout)
            ferr.seek(0)
            cerr = ferr.read()
            print(cerr)
        return asyrc

# Function for generating figures for benchmark output    
def benchfigs(rundims, shortrun):
    figs = []
    # FFT directions
    forwards = -1
    backwards = 1
    
    if 1 in rundims:
        dimension = 1

        nbatch = 1

        min1d = 256 if shortrun else 1024
        max1d = 4000 if shortrun else 536870912

        for inplace in [True, False]:
            fig = figure("1d_c2c" + ("inplace" if inplace else "outofplace"),
                         "1D complex transforms " + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3]:
                fig.runs.append( rundata("radix " + str(radix),
                                         dimension, nextpow(min1d, radix), max1d, nbatch,
                                         radix, [], "c2c", forwards, inplace) )
            figs.append(fig)

        for inplace in [True, False]:
            fig = figure("1d_r2c" + ("inplace" if inplace else "outofplace")
                         , "1D real-to-complex transforms " \
                         + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3]:
                fig.runs.append( rundata("radix " + str(radix),
                                         dimension, nextpow(min1d, radix), max1d, nbatch,
                                         radix, [], "r2c", forwards, inplace) )
            figs.append(fig)

            
        for inplace in [True, False]:
            fig = figure("1d_c2r" + ("inplace" if inplace else "outofplace"),
                         "1D complex-to-real transforms " \
                         + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3]:
                fig.runs.append( rundata("radix " + str(radix) ,
                                         dimension, nextpow(min1d, radix), max1d, nbatch,
                                         radix, [], "r2c", backwards, inplace) )
            figs.append(fig)


    if 2 in rundims:
        dimension = 2

        nbatch = 1
        min2d = 64 if shortrun else 128
        max2d = 8192 if shortrun else 32768

        for inplace in [True, False]:
            fig = figure("2d_c2c" + ("inplace" if inplace else "outofplace"),
                         "2D complex transforms " + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3]:
                fig.runs.append( rundata("radix " + str(radix), dimension,
                                         nextpow(min2d, radix), max2d, nbatch, radix, [1],
                                         "c2c",
                                         forwards, inplace) )
            figs.append(fig)

        for inplace in [True, False]:
            fig = figure("2d_r2c" + ("inplace" if inplace else "outofplace"),
                         "2D real-to-complex transforms " \
                         + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3]:
                fig.runs.append( rundata("radix " + str(radix), dimension,
                                         nextpow(min2d, radix), max2d, nbatch, radix, [1],
                                         "r2c",
                                         forwards, inplace) )
            figs.append(fig)

        for inplace in [True, False]:
            fig = figure("2d_c2r" + ("inplace" if inplace else "outofplace"),
                         "2D complex-to-real transforms " \
                         + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3]:
                fig.runs.append( rundata("radix " + str(radix), dimension,
                                         nextpow(min2d, radix), max2d, nbatch, radix, [1],
                                         "r2c",
                                         backwards, inplace) )
            figs.append(fig)

    if 3 in rundims:
        dimension = 3
        min3d = 16
        max3d = 128 if shortrun else 1024
        nbatch = 1
        
        for inplace in [True]:
            fig = figure("3d_c2c" + ("inplace" if inplace else "outofplace"),
                         "3D complex transforms " + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3, 5]:
                fig.runs.append( rundata("radix " + str(radix), dimension,
                                         nextpow(min3d, radix), max3d, nbatch, radix, [1,1],
                                         "c2c",
                                         forwards, inplace) )
            figs.append(fig)

        for inplace in [True, False]:
            fig = figure("3d_r2c" + ("inplace" if inplace else "outofplace")
                         , "3D real-to-complex transforms " \
                         + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3]:
                fig.runs.append( rundata("radix " + str(radix), dimension,
                                         nextpow(min3d, radix), max3d, nbatch, radix, [1,1],
                                         "r2c",
                                         forwards, inplace) )
            figs.append(fig)

        for inplace in [True, False]:
            fig = figure("3d_c2r" + ("inplace" if inplace else "outofplace"),
                         "3D complex-to-real transforms " \
                         + ("in-place" if inplace else "out-of-place"))
            for radix in [2, 3]:
                fig.runs.append( rundata("radix " + str(radix), dimension,
                                         nextpow(min3d, radix), max3d, nbatch, radix, [1,1],
                                         "r2c",
                                         backwards, inplace) )
            figs.append(fig)

    return figs

def efficiencyfigs(rundims, shortrun):
    figs = []
    
    # FFT directions
    forwards = -1
    backwards = 1

    inplace = True
    dimension = 1

    radix = 2

    min1d = 1024
    max1d = 1048576 if shortrun else 268435456 #pow(2,28) gives a floating type :(
    nbatch = 1
    while max1d > min1d:
        fig = figure("1d_c2c_batch" + str(nbatch) + "_radix" + str(radix),
                     "1D complex transforms " + ("in-place" if inplace else "out-of-place") + " radix " + str(radix) + " batch " + str(nbatch) )

        fig.runs.append( rundata("radix " + str(radix),
                                 dimension, nextpow(min1d, radix), max1d, nbatch,
                                 radix, [], "c2c", forwards, inplace) )
        figs.append(fig)
        nbatch *= 2
        max1d //= 2
        min1d //= 2
        min1d = max(min1d, 2^5)
    return figs

# Function for generating figures for a performance report
def reportfigs(rundims, shortrun):
    figs = []
    
    # FFT directions
    forwards = -1
    backwards = 1

    inplace = True
    
    if 1 in rundims:
        dimension = 1

        for min1d, max1d, nbatch in [[1024,536870912,1], [8,32768,100000]]:

            for radix in [2, 3, 5, 7]:
                fig = figure("1d_c2c" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "1D complex transforms with radix " + str(radix)\
                             + " and batch size " + str(nbatch) + "." )

                fig.runs.append( rundata("radix " + str(radix),
                                         dimension, nextpow(min1d, radix),
                                         max1d, nbatch,
                                         radix, [], "c2c", forwards,
                                         inplace) )
                figs.append(fig)

            for radix in [2, 3, 5, 7]:
                fig = figure("1d_r2c"\
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "1D real-to-complex transforms with radix "\
                             + str(radix) \
                             + " and batch size " + str(nbatch) + ".")
                fig.runs.append( rundata("radix " + str(radix),
                                         dimension, nextpow(min1d, radix),
                                         max1d, nbatch,
                                         radix, [], "r2c", forwards,
                                         inplace) )
                figs.append(fig)

            for radix in [2, 3, 5, 7]:
                fig = figure("1d_c2r" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "1D complex-to-real transforms with radix " \
                             + str(radix) \
                             + " and batch size " + str(nbatch) + "." )
                fig.runs.append( rundata("radix " + str(radix),
                                         dimension, nextpow(min1d, radix),
                                         max1d, nbatch,
                                         radix, [], "r2c", backwards,
                                         inplace) )
                figs.append(fig)

    if 2 in rundims:
        dimension = 2

        for min2d, max2d, nbatch in [[128,32768,1], [64,8192,100]]:

            for radix in [2, 3, 5]:
                fig = figure("2d_c2c" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch) ,
                             "2D complex transforms with radix " + str(radix)\
                             + " and batch size " + str(nbatch) + ".")
                fig.runs.append( rundata( "radix "+ str(radix),
                                         dimension,
                                         nextpow(min2d, radix), max2d,
                                         nbatch,
                                         radix, [1], "c2c",
                                         forwards, inplace) )
                figs.append(fig)

            for radix in [2, 3, 5]:
                fig = figure("2d_r2c" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "2D real-to-complex transforms with radix "\
                             + str(radix) \
                             + " and batch size " + str(nbatch) + ".")

                fig.runs.append( rundata( "radix " + str(radix),
                                         dimension,
                                         nextpow(min2d, radix), max2d,
                                         nbatch,
                                         radix, [1], "r2c",
                                         forwards, inplace) )
                figs.append(fig)

            for radix in [2, 3, 5]:
                fig = figure("2d_c2r" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "2D complex-to-real transforms with radix "\
                             + str(radix) +\
                             " and batch size " + str(nbatch) + ".")
                fig.runs.append( rundata("radix " + str(radix),
                                         dimension,
                                         nextpow(min2d, radix), max2d,
                                         nbatch,
                                         radix, [1], "r2c",
                                         backwards, inplace) )
                figs.append(fig)


            for radix in [2]:
                fig = figure("2d_c2c_r2" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "2D complex transforms "\
                             + "with aspect ratio N:2N with radix "\
                             + str(radix) + " and batch size " + str(nbatch) \
                             + ".")
                fig.runs.append( rundata( "radix 2",
                                         dimension, min2d, max2d, nbatch, 2,
                                         [2], "c2c",
                                         forwards, inplace) )
                figs.append(fig)

            for radix in [2]:
                fig = figure("2d_r2c_r2" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "2D real-to-complex transforms with radix "\
                             + str(radix) \
                             + " and batch size " + str(nbatch) + ".")
                fig.runs.append( rundata("radix 2",
                                         dimension, min2d, max2d, nbatch, 2,
                                         [2], "r2c",
                                         forwards, inplace) )
                figs.append(fig)

    if 3 in rundims:
        dimension = 3

        for min3d, max3d, nbatch in [[16,128,1],[4,64,100]]:
                
            for radix in [2, 3, 5]:
                fig = figure("3d_c2c" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "3D complex transforms with radix "\
                             + str(radix) \
                             + " and batch size " + str(nbatch) + ".")
                fig.runs.append( rundata("radix " + str(radix),
                                         dimension,
                                         nextpow(min3d, radix), max3d,
                                         nbatch,
                                         radix, [1,1], "c2c",
                                         forwards, inplace) )
                figs.append(fig)

            for radix in [2, 3]:
                fig = figure("3d_r2c" \
                             + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                             "3D real-to-complex transforms with radix "\
                             + str(radix)\
                             + " and batch size " + str(nbatch) + ".")

                fig.runs.append( rundata("radix " + str(radix),
                                         dimension,
                                         nextpow(min3d, radix), max3d,
                                         nbatch,
                                         radix, [1,1], "r2c",
                                         forwards, inplace) )
                figs.append(fig)


            fig = figure("3d_c2r" \
                         + "_radix" + str(radix) \
                         + "_batch" + str(nbatch),
                         "3D complex-to-real transforms with radix "\
                         + str(radix)
                         + " and batch size " + str(nbatch) + ".")
            for radix in [2]:
                fig.runs.append( rundata("radix " + str(radix),
                                         dimension,
                                         nextpow(min3d, radix), max3d,
                                         nbatch,
                                         radix, [1,1], "r2c",
                                         backwards, inplace) )
            figs.append(fig)

            fig = figure("3d_c2c_aspect" \
                         + "_radix" + str(radix) \
                         + "_batch" + str(nbatch),
                         "3D complex transforms "\
                         + "with aspect ratio N:N:16N with radix "\
                         + str(radix)\
                         + " and batch size " + str(nbatch) + ".")
            fig.runs.append( rundata("radix 2",
                                     dimension, min3d, max3d, nbatch, 2,
                                     [1,16], "c2c",
                                     forwards, inplace) )
            figs.append(fig)

            fig = figure("3d_r2c_aspect" \
                         + "_radix" + str(radix) \
                             + "_batch" + str(nbatch),
                         "3D real-to-complex transforms " \
                         + "with aspect ratio N:N:16N with radix " \
                         + str(radix)\
                         + " and batch size " + str(nbatch) + ".")
            fig.runs.append( rundata("radix 2",
                                     dimension, min3d, max3d, nbatch, 2,
                                     [1,16], "r2c",
                                     forwards, inplace) )
            figs.append(fig)
        
    return figs


def main(argv):
    dloadexe = None
    
    inlist = []
    outdirlist = []
    labellist = []

    docdir = "doc"
    
    dryrun = False
    nbatch = 1
    speedup = True
    datatype = "time"
    shortrun = False
    docformat = "pdf"
    devicenum = 0
    doAsy = True
    nsample = 10
    rundims = [1,2,3]
    runtype = "benchmark"
    secondtype = "none"
    
    try:
        opts, args = getopt.getopt(argv,"hb:D:f:Tt:i:o:l:S:sg:d:N:R:w:y:")
    except getopt.GetoptError:
        print("error in parsing arguments.")
        print(usage)
        sys.exit(2)
    for opt, arg in opts:
        if opt in ("-h"):
            print(usage)
            exit(0)
        elif opt in ("-b"):
            dloadexe = os.path.abspath(arg)
        elif opt in ("-i"):
            inlist.append(arg)
        elif opt in ("-o"):
            outdirlist.append(arg)
        elif opt in ("-l"):
            labellist.append(arg)
        elif opt in ("-w"):
            docdir = arg
        elif opt in ("-T"):
            dryrun = True
        elif opt in ("-s"):
            shortrun = True
        elif opt in ("-g"):
            if int(arg) == 0:
                doAsy = False
            if int(arg) == 1:
                doAsy = True
        elif opt in ("-d"):
            devicenum = int(arg)
        elif opt in ("-D"):
            rundims = []
            for val in arg.split(','):
                rundims.append(int(val))
        elif opt in ("-N"):
            nsample = int(arg)
        elif opt in ("-S"):
            if int(arg) == 0:
                speedup = False
            if int(arg) == 1:
                speedup = True
        elif opt in ("-t"):
            if arg not in ["time", "gflops", "roofline"]:
                print("data type must be time or gflops or roofline")
                print(usage)
                sys.exit(1)
            datatype = arg
        elif opt in ("-y"):
            if arg not in ["none", "gflops"]:
                print("data type must be gflops or none")
                print(usage)
                sys.exit(1)
            secondtype = arg
        elif opt in ("-R"):
            if arg not in ["report", "benchmark", "efficiency"]:
                print("data type must be gflops or none")
                print(usage)
                sys.exit(1)
            runtype = arg
            if runtype == "efficiency":
                datatype = "roofline"
        elif opt in ("-f"):
            goodvals = ["pdf", "docx"]
            if arg not in goodvals:
                print("error: format must in " + " ".join(goodvals))
                print(usage)
                sys.exit(1)
            docformat = arg

    print("rundims:")
    print(rundims)
    
    if not dryrun:
        if dloadexe != None:
            if not os.path.isfile(dloadexe):
                print("unable to find " + dloadexe)
                sys.exit(1)

        for i in inlist:
            if not os.path.isfile(i):
                print("unable to find " + i)
                print("please specify with -i")
                sys.exit(1)

    print("inputs:", inlist)
                
    if len(inlist) > len(labellist):
        for i in range(len(labellist), len(inlist)):
            labellist.append("dir" + str(i))
    print("run labels:", labellist)
    
    for idx in range(len(inlist)):
        inlist[idx] = os.path.abspath(inlist[idx])

    if len(inlist) > len(outdirlist):
        for i in range(len(outdirlist), len(inlist)):
            outdirlist.append(os.path.abspath("dir" + str(i)))

    for idx in range(len(outdirlist)):
        outdirlist[idx] = os.path.abspath(outdirlist[idx])
    print("data output directories:", outdirlist)

            
    if shortrun:
        print("short run")
    print("output format: " + docformat)
    print("device number: " + str(devicenum))

    docdir = os.path.abspath(docdir)
    
    print("document output in", docdir)
    if not os.path.exists(docdir):
        os.makedirs(docdir)
    
    for outdir in outdirlist:
        if not os.path.exists(outdir):
            os.makedirs(outdir)

            
    if not dryrun:
        import getspecs
        specs = "Host info:\n"
        specs += "\thostname: " + getspecs.gethostname() + "\n"
        specs += "\tcpu info: " + getspecs.getcpu() + "\n"
        specs += "\tram: " + getspecs.getram() + "\n"
        specs += "\tdistro: " + getspecs.getdistro() + "\n"
        specs += "\tkernel version: " + getspecs.getkernel() + "\n"
        specs += "\trocm version: " + getspecs.getrocmversion() + "\n"
        specs += "Device info:\n"
        specs += "\tdevice: " + getspecs.getdeviceinfo(devicenum) + "\n"
        specs += "\tvbios version: " + getspecs.getvbios(devicenum) + "\n"
        specs += "\tvram: " + getspecs.getvram(devicenum) + "\n"
        specs += "\tperformance level: " + getspecs.getperflevel(devicenum) + "\n"
        specs += "\tsystem clock: " + getspecs.getsclk(devicenum) + "\n"
        specs += "\tmemory clock: " + getspecs.getmclk(devicenum) + "\n"

        
        for outdir in outdirlist:
            with open(os.path.join(outdir, "specs.txt"), "w+") as f:
                f.write(specs)

            with open(os.path.join(outdir, "gpuid.txt"), "w") as f:
                f.write(getspecs.getgpuid(devicenum))

    figs = []

    if runtype == "benchmark":
        figs = benchfigs(rundims, shortrun)
    if runtype == "report":
        figs = reportfigs(rundims, shortrun)
    if runtype == "efficiency":
        figs = efficiencyfigs(rundims, shortrun)
    just1dc2crad2 = runtype == "efficiency"

    for idx, fig in enumerate(figs):
        for idx2, fig2 in enumerate(figs):
            if idx != idx2 and fig.name == fig2.name:
                print("figures have the same name!")
                print(fig.name)
                print(fig2.name)
                sys.exit(1)
    
    for fig in figs:
        print(fig.name)
        # Run the tests and put output in the outdirs:
        for run in fig.runs:
            if not dryrun:
                run.executerun(nsample, inlist, outdirlist, dloadexe)

        # Compile the data in the outdirs into figures in docdir:
        ncompare = len(inlist) if speedup else 0
        print(fig.labels(labellist))
        #plotgflops = runtype == "submission" and not datatype == "gflops"
        print(fig.asycmd(docdir, outdirlist, labellist, docformat, datatype, ncompare, secondtype, just1dc2crad2))
        fig.executeasy(docdir, outdirlist, labellist, docformat, datatype, ncompare, secondtype, just1dc2crad2)

    # Make the document in docdir:
    if docformat == "pdf":
        maketex(figs, docdir, outdirlist, labellist, nsample, secondtype)
    if docformat == "docx":
        makedocx(figs, docdir, nsample, secondtype)

    print("Finished!  Output in " + docdir)

def binaryisok(dirname, progname):
    prog = os.path.join(dirname, progname)
    return os.path.isfile(prog)

gflopstext = '''\
GFLOP/s are computed based on the Cooley--Tukey operation count \
for a radix-2 transform, and half that for in the case of \
real-complex transforms.  The rocFFT operation count may differ from \
this value: GFLOP/s is provided for the sake of comparison only.'''

# Function for generating a tex document in PDF format.
def maketex(figs, docdir, outdirlist, labellist, nsample, secondtype):
    
    header = '''\documentclass[12pt]{article}
\\usepackage{graphicx}
\\usepackage{url}
\\author{Malcolm Roberts}
\\begin{document}
'''
    texstring = header

    texstring += "\n\\section{Introduction}\n"
    
    texstring += "Each data point represents the median of " + str(nsample) + " values, with error bars showing the 95\\% confidence interval for the median.  All transforms are double-precision.\n\n"

    if secondtype == "gflops":
        texstring += gflopstext + "\n\n"

    
    texstring += "\\vspace{1cm}\n"
    
    # texstring += "\\begin{tabular}{ll}"
    # texstring += labelA +" &\\url{"+ dirA+"} \\\\\n"
    # if not dirB == None:
    #     texstring += labelB +" &\\url{"+ dirB+"} \\\\\n"
    # texstring += "\\end{tabular}\n\n"

    # texstring += "\\vspace{1cm}\n"
    
    texstring += "\n\\section{Device Specification}\n"
    for idx in range(len(outdirlist)):
        texstring += "\n\\subsection{" + labellist[idx]  + "}\n"
        specfilename = os.path.join(outdirlist[idx], "specs.txt")
        if os.path.isfile(specfilename):
            specs = ""
            with open(specfilename, "r") as f:
                specs = f.read()

            for line in specs.split("\n"):
                if line.startswith("Host info"):
                    texstring += "\\noindent " + line
                    texstring += "\\begin{itemize}\n"
                elif line.startswith("Device info"):
                    texstring += "\\end{itemize}\n"
                    texstring += line 
                    texstring += "\\begin{itemize}\n"
                else:
                    if line.strip() != "":
                        texstring += "\\item " + line + "\n"
            texstring += "\\end{itemize}\n"
            texstring += "\n"
        
    texstring += "\\clearpage\n"

    texstring += "\n\\section{Figures}\n"
    
    for idx, fig in enumerate(figs):
        print(fig.filename(docdir, "pdf"))
        print(fig.caption)
        texstring += '''
\\centering
\\begin{figure}[htbp]
   \\includegraphics[width=\\textwidth]{'''
        texstring += fig.filename("", "pdf")
        texstring += '''}
   \\caption{''' + fig.caption + '''}
\\end{figure}
'''
        if (idx % 2) == 0:
            texstring += "\\clearpage\n"
            
    texstring += "\n\\end{document}\n"
   
    fname = os.path.join(docdir, 'figs.tex')

    with open(fname, 'w') as outfile:
        outfile.write(texstring)

    fout = open(os.path.join(docdir, "texcmd.log"), 'w+')
    ferr = open(os.path.join(docdir, "texcmd.err"), 'w+')
                    
    latexcmd = ["latexmk", "-pdf", 'figs.tex']
    print(" ".join(latexcmd))
    texproc =  subprocess.Popen(latexcmd, cwd=docdir, stdout=fout, stderr=ferr,
                                env=os.environ.copy())
    texproc.wait()
    fout.close()
    ferr.close()
    texrc = texproc.returncode
    if texrc != 0:
        print("****tex fail****")

# Confert a PDF to an EMF using pdf2svg and inkscape.        
def pdf2emf(pdfname):
    svgname = pdfname.replace(".pdf",".svg")
    cmd_pdf2svg = ["pdf2svg", pdfname, svgname]
    proc = subprocess.Popen(cmd_pdf2svg, env=os.environ.copy())
    proc.wait()
    if proc.returncode != 0:
        print("pdf2svg failed!")
        sys.exit(1)

    emfname = pdfname.replace(".pdf",".emf")
    cmd_svg2emf = ["inkscape", svgname, "-M", emfname]
    proc = subprocess.Popen(cmd_svg2emf, env=os.environ.copy())
    proc.wait()
    if proc.returncode != 0:
        print("svg2emf failed!")
        sys.exit(1)
    
    return emfname

# Function for generating a docx using emf files and the docx package.
def makedocx(figs, outdir, nsample, secondtype):
    import docx

    document = docx.Document()

    document.add_heading('rocFFT benchmarks', 0)

    document.add_paragraph("Each data point represents the median of " + str(nsample) + " values, with error bars showing the 95% confidence interval for the median.  Transforms are double-precision, forward, and in-place.")

    if secondtype == "gflops":
        document.add_paragraph(gflopstext)
                               
    specfilename = os.path.join(outdir, "specs.txt")
    if os.path.isfile(specfilename):
        with open(specfilename, "r") as f:
            specs = f.read()
        for line in specs.split("\n"):
            document.add_paragraph(line)

    for fig in figs:
        print(fig.filename(outdir, "docx"))
        print(fig.caption)
        emfname = pdf2emf(fig.filename(outdir, "docx"))
        document.add_picture(emfname, width=docx.shared.Inches(6))
        document.add_paragraph(fig.caption)
                         
    document.save(os.path.join(outdir,'figs.docx'))
    
if __name__ == "__main__":
    main(sys.argv[1:])
                        
